import docx

doc1=docx.Document('D:\\application\\Odoo13\\template.docx')
# for i in range(0, len(doc1.paragraphs)):
#     text = doc1.paragraphs[i].text
#     print(text)

# doc2 = docx.Document()
# for i in range(0, len(doc1.paragraphs)):
#     text = doc1.paragraphs[i].text
#     text = text.replace('{{er_name}}', 'CIEE', 1)
#     text = text.replace('{{code}}', 'abc-123', 1)
#     text = text.replace('{{name}}', 'Tom 张三', 1)
#
#     doc2.add_paragraph(text, doc1.paragraphs[i].style)
#     # print(text)
# doc2.save('D:\\application\\Odoo13\\template-123.docx')

for paragraph in doc1.paragraphs:
    for run in paragraph.runs:
        # if "唐诗" in run.text:
        #     run.text = run.text.replace('唐诗','宋词')
        run.text = run.text.replace('{{er_name}}', 'CIEE', 1)
        run.text = run.text.replace('{{code}}', 'abc-123', 1)
        run.text = run.text.replace('{{name}}', 'Tom 张三', 1)

doc1.save('template-123.docx')
